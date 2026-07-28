#!/usr/bin/env python3
"""Batch-generate BD meeting memos from meeting notes + web-search research evidence.

Pairs each meeting-notes .docx with its matching NN_research_evidence.md file
(matched by two-digit prefix), builds a prompt packet from the bd-market-memo
skill instructions, and either:
  --dry-run : writes the assembled prompt packet as the "memo" output, with no
              API call, so an instructor can review exactly what would be sent.
  (live)    : calls the Anthropic Messages API to draft the memo.

Always writes, per meeting: <out-dir>/<NN>_<slug>_memo.md and
<out-dir>/<NN>_<slug>_audit.json. Continues past a single meeting's failure;
exits non-zero if any meeting failed.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
import time
import traceback
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

try:
    import docx  # python-docx
except ImportError:  # pragma: no cover - guarded at runtime, not at import time
    docx = None

DEFAULT_MODEL = "claude-sonnet-5"
MEETING_PREFIX_RE = re.compile(r"^(\d{2})[_-]")
SOURCE_HEADER_RE = re.compile(r"^###\s*(\d+)?\.?\s*(.+?)\s*$")
FIELD_RE = re.compile(r"^-\s*\*\*(.+?):\*\*\s*(.*)$")


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass
class Source:
    title: str
    publisher: str = ""
    date: str = ""
    url: str = ""
    relevance_score: Optional[int] = None
    human_review_flag: str = "None"
    meeting_issue_connection: str = ""

    @property
    def missing_fields(self) -> list[str]:
        missing = []
        if not self.title:
            missing.append("title")
        if not self.publisher:
            missing.append("publisher")
        if not self.date:
            missing.append("date")
        if not self.url:
            missing.append("url")
        return missing

    @property
    def eligible(self) -> bool:
        return self.relevance_score == 3 and not self.missing_fields

    def to_dict(self) -> dict:
        return {
            "title": self.title,
            "publisher": self.publisher,
            "date": self.date,
            "url": self.url,
            "relevance_score": self.relevance_score,
            "human_review_flag": self.human_review_flag,
            "meeting_issue_connection": self.meeting_issue_connection,
            "missing_fields": self.missing_fields,
            "eligible": self.eligible,
        }


@dataclass
class MeetingResult:
    meeting_number: str
    meeting_file: str
    research_file: str
    status: str = "ok"  # ok | insufficient_input | error
    mode: str = "dry_run"
    model: str = DEFAULT_MODEL
    elapsed_seconds: float = 0.0
    source_count: int = 0
    eligible_source_count: int = 0
    sources: list = field(default_factory=list)
    missing_fields: list = field(default_factory=list)
    relevance_failures: list = field(default_factory=list)
    human_review_flags: list = field(default_factory=list)
    memo_path: str = ""
    error: Optional[str] = None

    def to_dict(self) -> dict:
        return {
            "meeting_number": self.meeting_number,
            "meeting_file": self.meeting_file,
            "research_file": self.research_file,
            "status": self.status,
            "mode": self.mode,
            "model": self.model,
            "elapsed_seconds": round(self.elapsed_seconds, 3),
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "source_count": self.source_count,
            "eligible_source_count": self.eligible_source_count,
            "sources": self.sources,
            "missing_fields": self.missing_fields,
            "relevance_failures": self.relevance_failures,
            "human_review_flags": self.human_review_flags,
            "memo_path": self.memo_path,
            "error": self.error,
        }


# ---------------------------------------------------------------------------
# Input discovery / matching
# ---------------------------------------------------------------------------

def find_prefix(filename: str) -> Optional[str]:
    m = MEETING_PREFIX_RE.match(filename)
    return m.group(1) if m else None


def discover_meetings(notes_dir: Path, research_dir: Path) -> list[dict]:
    """Pair notes .docx files with research evidence .md files by 2-digit prefix."""
    notes_by_prefix: dict[str, Path] = {}
    for p in sorted(notes_dir.glob("*.docx")):
        prefix = find_prefix(p.name)
        if prefix:
            notes_by_prefix[prefix] = p

    research_by_prefix: dict[str, Path] = {}
    for p in sorted(research_dir.glob("*_research_evidence.md")):
        prefix = find_prefix(p.name)
        if prefix:
            research_by_prefix[prefix] = p

    all_prefixes = sorted(set(notes_by_prefix) | set(research_by_prefix))
    pairs = []
    for prefix in all_prefixes:
        pairs.append({
            "prefix": prefix,
            "notes_path": notes_by_prefix.get(prefix),
            "research_path": research_by_prefix.get(prefix),
        })
    return pairs


# ---------------------------------------------------------------------------
# Extraction
# ---------------------------------------------------------------------------

def extract_docx_text(path: Path) -> str:
    """Extract text in actual document order (paragraphs interleaved with tables).

    python-docx's `document.paragraphs` and `document.tables` are separate flat
    lists that don't preserve body order, so a docx with a metadata table (e.g.
    Date/Participants) followed by prose paragraphs would come out reordered if
    each list were walked separately. Walk the body XML directly instead.
    """
    if docx is None:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install -r requirements.txt"
        )
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = docx.Document(str(path))
    parts = []
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            text = Paragraph(child, document).text.strip()
            if text:
                parts.append(text)
        elif child.tag == qn("w:tbl"):
            table = Table(child, document)
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_research_evidence(markdown_text: str) -> list[Source]:
    """Parse the NN_research_evidence.md convention:

    ### <n>. <Title>
    - **Publisher:** ...
    - **Publication/update date:** ...
    - **URL:** ...
    - **Meeting issue it connects to:** ...
    - **Relevance score:** 3
    - **Human-review flag:** ...
    """
    sources: list[Source] = []
    current: Optional[Source] = None

    field_map = {
        "publisher": "publisher",
        "publication/update date": "date",
        "publication date": "date",
        "url": "url",
        "meeting issue it connects to": "meeting_issue_connection",
        "human-review flag": "human_review_flag",
    }

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()

        header = SOURCE_HEADER_RE.match(line)
        if header and line.startswith("###"):
            if current is not None:
                sources.append(current)
            title = header.group(2).strip().replace('"', "")
            current = Source(title=title)
            continue

        if current is None:
            continue

        field_match = FIELD_RE.match(line)
        if field_match:
            key = field_match.group(1).strip().lower()
            value = field_match.group(2).strip()
            if key == "relevance score":
                m = re.search(r"\d+", value)
                current.relevance_score = int(m.group()) if m else None
            elif key in field_map:
                setattr(current, field_map[key], value)

    if current is not None:
        sources.append(current)

    return sources


# ---------------------------------------------------------------------------
# Skill / prompt assembly
# ---------------------------------------------------------------------------

def load_skill_instructions(skill_dir: Path) -> str:
    skill_md = skill_dir / "SKILL.md"
    if not skill_md.exists():
        raise FileNotFoundError(f"Skill definition not found at {skill_md}")

    parts = [skill_md.read_text(encoding="utf-8")]
    references_dir = skill_dir / "references"
    if references_dir.exists():
        for ref in sorted(references_dir.glob("*.md")):
            parts.append(f"\n\n---\n# Reference: {ref.name}\n\n{ref.read_text(encoding='utf-8')}")
    return "\n".join(parts)


def build_user_prompt(meeting_number: str, notes_text: str, research_text: str) -> str:
    return f"""Draft the BD meeting memo for meeting {meeting_number} following the SKILL.md
instructions and the output-schema.md format exactly (Shape 1 if input is sufficient,
Shape 2 "Insufficient Input" if it is not).

## Meeting Notes (source of Facts about the meeting)

{notes_text}

## Research Evidence (source of Market & Competitive Evidence)

{research_text}

Apply Rule A (complete follow-ups, TBD for missing owner/timeline) and Rule B
(score-3 relevance gate) exactly as specified. Do not send messages, modify CRM
records, make commitments, or give legal advice."""


def build_prompt_packet(system_prompt: str, user_prompt: str, model: str) -> dict:
    return {
        "model": model,
        "system": system_prompt,
        "messages": [
            {"role": "user", "content": user_prompt},
        ],
    }


# ---------------------------------------------------------------------------
# Per-meeting processing
# ---------------------------------------------------------------------------

def render_dry_run_memo(prefix: str, packet: dict) -> str:
    lines = [
        f"# DRY RUN — Prompt Packet for Meeting {prefix}",
        "",
        "No API call was made. This file is the complete prompt that would be",
        "sent to the Anthropic Messages API in live mode.",
        "",
        f"**Model:** {packet['model']}",
        "",
        "## System Prompt",
        "",
        "```markdown",
        packet["system"],
        "```",
        "",
        "## User Message",
        "",
        "```markdown",
        packet["messages"][0]["content"],
        "```",
    ]
    return "\n".join(lines)


def call_anthropic(packet: dict) -> str:
    import anthropic  # imported lazily so --dry-run never requires the package

    client = anthropic.Anthropic()  # reads ANTHROPIC_API_KEY from environment
    response = client.messages.create(
        model=packet["model"],
        max_tokens=4096,
        system=packet["system"],
        messages=packet["messages"],
    )
    return "".join(
        block.text for block in response.content if getattr(block, "type", None) == "text"
    )


def slugify(meeting_file: Optional[Path], prefix: str) -> str:
    if meeting_file is None:
        return f"meeting_{prefix}"
    stem = meeting_file.stem
    stem = re.sub(rf"^{prefix}[_-]?", "", stem)
    return stem or f"meeting_{prefix}"


def process_meeting(
    pair: dict,
    skill_instructions: str,
    out_dir: Path,
    model: str,
    dry_run: bool,
) -> MeetingResult:
    prefix = pair["prefix"]
    notes_path: Optional[Path] = pair["notes_path"]
    research_path: Optional[Path] = pair["research_path"]

    result = MeetingResult(
        meeting_number=prefix,
        meeting_file=notes_path.name if notes_path else "MISSING",
        research_file=research_path.name if research_path else "MISSING",
        mode="dry_run" if dry_run else "live",
        model=model,
    )

    start = time.monotonic()
    slug = slugify(notes_path, prefix)
    memo_path = out_dir / f"{prefix}_{slug}_memo.md"
    audit_path = out_dir / f"{prefix}_{slug}_audit.json"
    result.memo_path = str(memo_path)

    try:
        if notes_path is None or research_path is None:
            result.status = "insufficient_input"
            missing = []
            if notes_path is None:
                missing.append("meeting notes file not found for this prefix")
            if research_path is None:
                missing.append("research evidence file not found for this prefix")
            result.missing_fields = missing
            memo_text = render_insufficient_input(
                prefix,
                result,
                missing_reason="\n".join(missing),
                checks={
                    "Meeting notes file present and non-empty": notes_path is not None,
                    "Research evidence file present and non-empty": research_path is not None,
                    "At least one source meets the relevance-3 bar": False,
                },
            )
            memo_path.write_text(memo_text, encoding="utf-8")
            result.elapsed_seconds = time.monotonic() - start
            audit_path.write_text(json.dumps(result.to_dict(), indent=2), encoding="utf-8")
            return result

        notes_text = extract_docx_text(notes_path)
        research_text = research_path.read_text(encoding="utf-8")

        if not notes_text.strip() or not research_text.strip():
            result.status = "insufficient_input"
            missing = []
            if not notes_text.strip():
                missing.append("meeting notes file is empty")
            if not research_text.strip():
                missing.append("research evidence file is empty")
            result.missing_fields = missing
            memo_text = render_insufficient_input(
                prefix,
                result,
                missing_reason="\n".join(missing),
                checks={
                    "Meeting notes file present and non-empty": bool(notes_text.strip()),
                    "Research evidence file present and non-empty": bool(research_text.strip()),
                    "At least one source meets the relevance-3 bar": False,
                },
            )
            memo_path.write_text(memo_text, encoding="utf-8")
            result.elapsed_seconds = time.monotonic() - start
            audit_path.write_text(json.dumps(result.to_dict(), indent=2), encoding="utf-8")
            return result

        sources = parse_research_evidence(research_text)
        result.source_count = len(sources)
        result.sources = [s.to_dict() for s in sources]
        result.missing_fields = [
            f"{s.title}: missing {', '.join(s.missing_fields)}"
            for s in sources
            if s.missing_fields
        ]
        result.relevance_failures = [
            f"{s.title} (score={s.relevance_score})"
            for s in sources
            if s.relevance_score != 3
        ]
        flag_set: set[str] = set()
        for s in sources:
            if not s.eligible or not s.human_review_flag:
                continue
            for token in re.split(r"[;,]", s.human_review_flag):
                token = token.strip()
                if token and token.lower() != "none":
                    flag_set.add(token)
        result.human_review_flags = sorted(flag_set)
        eligible = [s for s in sources if s.eligible]
        result.eligible_source_count = len(eligible)

        if not eligible:
            result.status = "insufficient_input"
            memo_text = render_insufficient_input(
                prefix,
                result,
                missing_reason=(
                    f"{len(sources)} candidate source(s) were supplied; none met the "
                    "relevance-3 bar required by Rule B, so a memo cannot be drafted."
                ),
                checks={
                    "Meeting notes file present and non-empty": True,
                    "Research evidence file present and non-empty": True,
                    "At least one source meets the relevance-3 bar": False,
                },
            )
            memo_path.write_text(memo_text, encoding="utf-8")
            result.elapsed_seconds = time.monotonic() - start
            audit_path.write_text(json.dumps(result.to_dict(), indent=2), encoding="utf-8")
            return result

        system_prompt = skill_instructions
        user_prompt = build_user_prompt(prefix, notes_text, research_text)
        packet = build_prompt_packet(system_prompt, user_prompt, model)

        if dry_run:
            memo_text = render_dry_run_memo(prefix, packet)
        else:
            memo_text = call_anthropic(packet)

        memo_path.write_text(memo_text, encoding="utf-8")
        result.status = "ok"
        result.elapsed_seconds = time.monotonic() - start
        audit_path.write_text(json.dumps(result.to_dict(), indent=2), encoding="utf-8")
        return result

    except Exception as exc:  # noqa: BLE001 - isolate failure to this meeting
        result.status = "error"
        result.error = f"{type(exc).__name__}: {exc}"
        result.elapsed_seconds = time.monotonic() - start
        try:
            audit_path.write_text(json.dumps(result.to_dict(), indent=2), encoding="utf-8")
        except Exception:
            pass
        return result


def render_insufficient_input(
    prefix: str,
    result: MeetingResult,
    missing_reason: str,
    checks: Optional[dict[str, bool]] = None,
) -> str:
    checks = checks or {}
    default_checks = {
        "Meeting notes file present and non-empty": False,
        "Research evidence file present and non-empty": False,
        "At least one source meets the relevance-3 bar": False,
    }
    default_checks.update(checks)
    checklist = "\n".join(
        f"- [{'x' if ok else ' '}] {label}" for label, ok in default_checks.items()
    )

    return f"""# BD Meeting Memo — Status Report

**Meeting file:** {result.meeting_file}
**Research evidence file:** {result.research_file}
**Memo generated:** {datetime.now(timezone.utc).isoformat()}
**Status:** Insufficient Input — memo not drafted

## What was checked

{checklist}

## What is missing

{missing_reason}

## What a human needs to supply or fix

- Confirm the meeting notes and research evidence files for meeting {prefix} both exist and are non-empty.
- Confirm the research evidence file lists title/publisher/date/URL and a relevance score for each candidate source.
- If all sources scored below 3, re-run research with a more targeted query before drafting a memo.

No memo body, follow-up table, or evidence section is produced for this meeting.
"""


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args(argv: Optional[list[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--notes-dir", required=True, type=Path, help="Directory of NN_*.docx meeting notes")
    parser.add_argument("--research-dir", required=True, type=Path, help="Directory of NN_research_evidence.md files")
    parser.add_argument("--out-dir", required=True, type=Path, help="Directory to write memos and audit records")
    parser.add_argument("--model", default=DEFAULT_MODEL, help=f"Anthropic model id (default: {DEFAULT_MODEL})")
    parser.add_argument("--dry-run", action="store_true", help="Assemble prompt packets, make no API calls")
    parser.add_argument(
        "--skill-dir",
        type=Path,
        default=None,
        help="Path to the bd-market-memo skill directory (default: <repo>/.claude/skills/bd-market-memo)",
    )
    return parser.parse_args(argv)


def main(argv: Optional[list[str]] = None) -> int:
    args = parse_args(argv)

    notes_dir: Path = args.notes_dir
    research_dir: Path = args.research_dir
    out_dir: Path = args.out_dir
    skill_dir: Path = args.skill_dir or (Path(__file__).resolve().parent.parent / ".claude" / "skills" / "bd-market-memo")

    if not notes_dir.is_dir():
        print(f"ERROR: --notes-dir does not exist or is not a directory: {notes_dir}", file=sys.stderr)
        return 2
    if not research_dir.is_dir():
        print(f"ERROR: --research-dir does not exist or is not a directory: {research_dir}", file=sys.stderr)
        return 2

    out_dir.mkdir(parents=True, exist_ok=True)

    if not args.dry_run:
        import os

        if not os.environ.get("ANTHROPIC_API_KEY"):
            print(
                "ERROR: ANTHROPIC_API_KEY is not set in the environment. "
                "Live mode requires it. Refusing to proceed.",
                file=sys.stderr,
            )
            return 2

    try:
        skill_instructions = load_skill_instructions(skill_dir)
    except FileNotFoundError as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2

    pairs = discover_meetings(notes_dir, research_dir)
    if not pairs:
        print(f"No meetings found matching a two-digit prefix in {notes_dir} / {research_dir}", file=sys.stderr)
        return 2

    print(f"Found {len(pairs)} meeting(s). Mode: {'DRY RUN' if args.dry_run else 'LIVE'}. Model: {args.model}")

    results: list[MeetingResult] = []
    for pair in pairs:
        prefix = pair["prefix"]
        print(f"[{prefix}] processing...", end=" ")
        try:
            result = process_meeting(pair, skill_instructions, out_dir, args.model, args.dry_run)
        except Exception as exc:  # noqa: BLE001 - last-resort safety net per meeting
            result = MeetingResult(
                meeting_number=prefix,
                meeting_file=pair["notes_path"].name if pair["notes_path"] else "MISSING",
                research_file=pair["research_path"].name if pair["research_path"] else "MISSING",
                mode="dry_run" if args.dry_run else "live",
                model=args.model,
                status="error",
                error=f"{type(exc).__name__}: {exc}",
            )
            traceback.print_exc(file=sys.stderr)
        results.append(result)
        print(result.status)

    failed = [r for r in results if r.status == "error"]
    insufficient = [r for r in results if r.status == "insufficient_input"]
    ok = [r for r in results if r.status == "ok"]

    print()
    print(f"Summary: {len(ok)} ok, {len(insufficient)} insufficient-input, {len(failed)} failed (of {len(results)})")
    if failed:
        print("Failed meetings:", ", ".join(r.meeting_number for r in failed), file=sys.stderr)

    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
